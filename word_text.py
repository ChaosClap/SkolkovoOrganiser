from docx import Document

def extract_text_from_paragraphs(paragraphs):
    return [p.text for p in paragraphs if p.text.strip()]

def extract_text_from_table(table):
    texts = []
    for row in table.rows:
        for cell in row.cells:
            # Avoid duplicate cells due to shared cell objects
            cell_texts = extract_text_from_paragraphs(cell.paragraphs)
            texts.extend(cell_texts)
            # Recursively check nested tables
            for nested_table in cell.tables:
                texts.extend(extract_text_from_table(nested_table))
    return texts

def extract_all_text(docx_path):
    document = Document(docx_path)
    all_text = []

    # Основной текст документа
    all_text.extend(extract_text_from_paragraphs(document.paragraphs))

    # Текст из всех таблиц
    for table in document.tables:
        all_text.extend(extract_text_from_table(table))

    # Верхние и нижние колонтитулы
    for section in document.sections:
        header = section.header
        footer = section.footer
        all_text.extend(extract_text_from_paragraphs(header.paragraphs))
        all_text.extend(extract_text_from_paragraphs(footer.paragraphs))
        for table in header.tables:
            all_text.extend(extract_text_from_table(table))
        for table in footer.tables:
            all_text.extend(extract_text_from_table(table))

    return all_text

# Пример использования
docx_file = "startup_template.docx"
text_content = extract_all_text(docx_file)

print("\n--- Найденный текст ---")
# Удаляем повторы, сохраняя порядок
seen = set()
unique_text = []
for line in text_content:
    stripped_line = line.strip()
    if stripped_line and stripped_line not in seen:
        unique_text.append(stripped_line)
        seen.add(stripped_line)

print("\n--- Найденный уникальный текст ---")
for line in unique_text:
    print(line)

