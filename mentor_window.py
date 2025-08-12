import sys
import requests
import io
from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QLabel, QScrollArea, QTextEdit, QPushButton
)
from PyQt5.QtCore import Qt
from docx import Document
from datetime import datetime
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PyQt5.QtWidgets import (
    QMainWindow, QWidget, QVBoxLayout, QLabel, QPushButton,
    QListWidget, QMessageBox, QHBoxLayout, QTextEdit
)

import subprocess  # для открытия проводника
from pathlib import Path

from PyQt5.QtGui import QPixmap
from io import BytesIO
from PyQt5.QtWidgets import QFileDialog

from profile_utils import (
    save_profile_locally, load_profile_locally, save_avatar_locally,
    get_avatar_path, get_avatar_bytes, debug_print_all_cells
)

import os
import json
from PIL import Image



API_URL = "http://127.0.0.1:8000/api/startups" # Подставь IP сервера, когда будет нужно

all_placeholders = [
    "date", "company", "industry", "description", "website", "technology", "sales", "presentation",
    "goals", "investment", "status", "plan",  # <- эти могли отсутствовать
    "full_name", "position", "experience", "skills", "achievements"
]


class MentorWindow(QMainWindow):
    def __init__(self, username, token):
        super().__init__()
        self.token = token
        self.username = username
        self.setWindowTitle("Интерфейс ментора")
        self.resize(1200, 600)

        self.current_page = 0
        self.page_size = 20
        self.total_items = 0

        self.prev_button = QPushButton("← Назад")
        self.next_button = QPushButton("Вперёд →")

        self.prev_button.clicked.connect(self.prev_page)
        self.next_button.clicked.connect(self.next_page)

        self.profile_data = self.get_profile_info()

        save_profile_locally(self.profile_data)
        save_avatar_locally(self.token)

        # Левая часть — аватар и профиль
        self.avatar_label = QLabel()
        self.avatar_label.setAlignment(Qt.AlignCenter)
        self.load_avatar()

        self.upload_avatar_button = QPushButton("Загрузить аватар")
        self.upload_avatar_button.clicked.connect(self.upload_avatar)

        self.profile_text = QTextEdit()
        self.profile_text.setReadOnly(True)
        self.profile_text.setMinimumWidth(300)
        self.profile_text.setStyleSheet("background-color: #f0f0f0; padding: 10px;")
        self.set_profile_info()

        left_panel = QVBoxLayout()
        avatar_container = QHBoxLayout()
        avatar_container.addStretch()
        avatar_container.addWidget(self.avatar_label)
        avatar_container.addStretch()
        left_panel.addLayout(avatar_container)
        left_panel.addWidget(self.upload_avatar_button)
        left_panel.addWidget(self.profile_text)

        # Центр — список стартапов
        self.list_widget = QListWidget()
        self.list_widget.currentRowChanged.connect(self.display_startup_details)

        self.export_button = QPushButton("Сформировать отчёт")
        self.status_label = QLabel("")

        center_panel = QVBoxLayout()
        center_panel.addWidget(QLabel("Список стартапов"))
        center_panel.addWidget(self.list_widget)
        center_panel.addWidget(self.export_button)
        center_panel.addWidget(self.status_label)
        nav_layout = QHBoxLayout()
        nav_layout.addWidget(self.prev_button)
        nav_layout.addWidget(self.next_button)
        center_panel.addLayout(nav_layout)

        # Правая часть — аватар и детали стартапа
        self.startup_avatar_label = QLabel()
        self.startup_avatar_label.setAlignment(Qt.AlignCenter)
        self.startup_avatar_label.setFixedSize(128, 128)
        self.startup_avatar_label.setStyleSheet("border: 1px solid gray; background-color: #eee;")

        self.detail_text = QTextEdit()
        self.detail_text.setReadOnly(True)
        self.detail_text.setMinimumWidth(400)
        self.detail_text.setStyleSheet("background-color: #f9f9f9; padding: 10px;")

        right_panel = QVBoxLayout()
        right_panel.addWidget(QLabel("Информация о стартапе"))
        right_panel.addWidget(self.startup_avatar_label)
        right_panel.addWidget(self.detail_text)

        # Главная раскладка
        main_layout = QHBoxLayout()
        main_layout.addLayout(left_panel)
        main_layout.addLayout(center_panel)
        main_layout.addLayout(right_panel)

        container = QWidget()
        container.setLayout(main_layout)
        self.setCentralWidget(container)

        self.startups = []
        self.export_button.clicked.connect(self.export_selected_startup)
        self.load_startups()

    def update_nav_buttons(self):
        self.prev_button.setEnabled(self.current_page > 0)
        max_pages = (self.total_items - 1) // self.page_size
        self.next_button.setEnabled(self.current_page < max_pages)

    def next_page(self):
        self.current_page += 1
        self.load_startups()

    def prev_page(self):
        if self.current_page > 0:
            self.current_page -= 1
            self.load_startups()

    def set_profile_info(self):
        profile = self.profile_data
        text = f"""
<b>ФИО:</b> {profile.get('full_name', '—')}<br>
<b>Должность:</b> {profile.get('position', '—')}<br><br>
<b>Опыт:</b><br>{profile.get('experience', '—')}<br><br>
<b>Компетенции:</b><br>{profile.get('skills', '—')}<br><br>
<b>Достижения:</b><br>{profile.get('achievements', '—')}
"""
        self.profile_text.setHtml(text)

    def display_startup_details(self, index):
        if index == -1:
            self.detail_text.clear()
            self.startup_avatar_label.clear()
            self.selected_startup = None
            return

        startup_id = self.startups[index]["id"]
        headers = {"Authorization": f"Bearer {self.token}"}

        try:
            response = requests.get(f"http://127.0.0.1:8000/api/startups/{startup_id}", headers=headers)
            if response.status_code == 200:
                data = response.json()
                self.selected_startup = data  # <- сохранить подробные данные здесь!
                details = f"""
    <b>Компания:</b> {data.get('company_name', '—')}<br>
    <b>Индустрия:</b> {data.get('industry', '—')}<br><br>
    <b>Описание:</b><br>{data.get('description', '—')}<br><br>
    <b>Сайт:</b> {data.get('website', '—')}<br><br>
    <b>Технология:</b><br>{data.get('technology', '—')}<br><br>
    <b>Продажи:</b><br>{data.get('sales', '—')}<br><br>
    <b>Презентация:</b><br>{data.get('presentation', '—')}<br><br>
    <b>Цели:</b><br>{data.get('goals', '—')}<br><br>
    <b>Инвестиции:</b><br>{data.get('investment', '—')}
    """
                self.detail_text.setHtml(details)
            else:
                self.detail_text.setHtml("⚠ Не удалось загрузить данные о стартапе")

            # Загружаем аватар
            r_avatar = requests.get(f"http://127.0.0.1:8000/api/startups/{startup_id}/avatar", headers=headers)
            if r_avatar.status_code == 200:
                pixmap = QPixmap()
                if pixmap.loadFromData(r_avatar.content):
                    self.startup_avatar_label.setPixmap(
                        pixmap.scaled(128, 128, Qt.KeepAspectRatio, Qt.SmoothTransformation))
                else:
                    self.startup_avatar_label.clear()
            else:
                self.startup_avatar_label.clear()

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка при получении деталей: {e}")

    def load_startups(self):
        try:
            headers = {"Authorization": f"Bearer {self.token}"}
            skip = self.current_page * self.page_size
            url = f"http://127.0.0.1:8000/api/startups?skip={skip}&limit={self.page_size}"
            response = requests.get(url, headers=headers)

            if response.status_code == 200:
                data = response.json()
                self.total_items = data["total"]
                self.startups = data["items"]

                self.list_widget.clear()
                for s in self.startups:
                    self.list_widget.addItem(s.get("company_name", "Без названия"))

                self.update_nav_buttons()

            else:
                QMessageBox.warning(self, "Ошибка", f"Код {response.status_code}: {response.text}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки: {e}")

    def export_selected_startup(self):
        if not self.selected_startup:
            QMessageBox.warning(self, "Выбор", "Выберите стартап для экспорта")
            return

        try:
            self.export_to_word(self.selected_startup)
            self.status_label.setText(f"✅ Отчёт создан: {self.selected_startup['company_name']}")
        except Exception as e:
            import traceback
            print("❌ Ошибка при создании отчёта:")
            traceback.print_exc()
            QMessageBox.critical(self, "Ошибка экспорта", str(e))

    def export_to_word(self, form_data):
        try:
            now = datetime.now()
            date_str = now.strftime("%d.%m.%Y %H:%M")

            form_data = {
                "company_name": form_data.get("company_name", ""),
                "company_ogrn": form_data.get("company_ogrn", ""),
                "representative_name": form_data.get("representative_name", ""),
                "email": form_data.get("email", ""),
                "phone": form_data.get("phone", ""),
                "industry": form_data.get("industry", ""),
                "description": form_data.get("description", ""),
                "website": form_data.get("website", ""),
                "technology": form_data.get("technology", ""),
                "sales": form_data.get("sales", ""),
                "presentation": form_data.get("presentation", ""),
                "goals": form_data.get("goals", ""),
                "investment": form_data.get("investment", ""),
                "current_stage": form_data.get("current_stage", ""),
                "six_month_plan": form_data.get("six_month_plan", ""),
                "date": date_str,
            }

            mentor_data = load_profile_locally()
            form_data["full_name"] = mentor_data.get("full_name", "")
            form_data["position"] = mentor_data.get("position", "")
            form_data["experience"] = mentor_data.get("experience", "")
            form_data["skills"] = mentor_data.get("skills", "")
            form_data["achievements"] = mentor_data.get("achievements", "")

            template_path = "startup_template.docx"

            # Папка для архива
            archive_dir = Path("diagnostics_archive")
            archive_dir.mkdir(exist_ok=True)

            # Уникальное имя
            base_name = f"{form_data['company_name']}_{now.strftime('%d.%m.%Y')}"
            output_path = archive_dir / f"{base_name}.docx"
            counter = 1
            while output_path.exists():
                output_path = archive_dir / f"{base_name}_{counter}.docx"
                counter += 1

            doc = Document(template_path)

            def replace_paragraph_content(paragraph):
                full_text = "".join(run.text for run in paragraph.runs)
                for key, value in form_data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in full_text:
                        for run in paragraph.runs:
                            run.text = ""
                        new_text = full_text.replace(placeholder, str(value))
                        run = paragraph.add_run(new_text)
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.name = "Arial"
                        run.font.italic = True
                        run.font.size = Pt(12)
                        break

            for paragraph in doc.paragraphs:
                replace_paragraph_content(paragraph)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_paragraph_content(paragraph)

            # Вставка аватара вместо 3-й картинки
            avatar_path = get_avatar_path()
            if avatar_path:
                image_count = 0
                replaced = False
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if run._element.xpath('.//w:drawing'):
                                        image_count += 1
                                        if image_count == 3:
                                            run.clear()
                                            run.add_picture(avatar_path, width=Inches(2))
                                            replaced = True
                                            print("✅ Аватар заменил 3-ю картинку в таблице")
                                            break
                                if replaced:
                                    break
                            if replaced:
                                break
                        if replaced:
                            break
                    if replaced:
                        break
                if not replaced:
                    print("⚠ 3-я картинка не найдена, аватар не вставлен.")

            doc.save(output_path)
            print(f"\n✅ Файл создан: {output_path}")
            self.status_label.setText(f"Файл создан: {output_path.name}")

            # Открываем проводник и выделяем файл
            try:
                subprocess.run(['explorer', '/select,', str(output_path.resolve())])
            except Exception as e:
                print(f"⚠ Не удалось открыть проводник: {e}")

        except Exception as e:
            import traceback
            error_msg = traceback.format_exc()
            QMessageBox.critical(self, "Ошибка экспорта", f"Произошла ошибка:\n{str(e)}\n\nПодробности:\n{error_msg}")

    def get_profile_info(self):
        headers = {"Authorization": f"Bearer {self.token}"}
        response = requests.get("http://127.0.0.1:8000/auth/me", headers=headers)
        if response.status_code == 200:
            return response.json()
        else:
            return {}

    def load_avatar(self):
        try:
            headers = {"Authorization": f"Bearer {self.token}"}
            response = requests.get("http://127.0.0.1:8000/auth/avatar", headers=headers)

            if response.status_code == 200:
                pixmap = QPixmap()
                if pixmap.loadFromData(response.content):
                    print("✅ Картинка успешно загружена и вставлена")
                    self.avatar_label.setPixmap(pixmap.scaled(128, 128, Qt.KeepAspectRatio, Qt.SmoothTransformation))
                else:
                    print("❌ QPixmap не смог загрузить изображение (формат не распознан)")
                    QMessageBox.warning(self, "Ошибка", "Не удалось распознать изображение")
            else:
                print(f"ℹ️ Сервер вернул статус: {response.status_code}")
            #    QMessageBox.warning(self, "Ошибка", f"Аватар не найден ({response.status_code})")

        except Exception as e:
            print("Ошибка при загрузке аватара:", e)
            QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки: {e}")

    def upload_avatar(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Выберите аватар", "", "Images (*.png *.jpg *.jpeg)")

        if file_path:
            try:
                img = Image.open(file_path)
                max_width, max_height = 360, 280
                original_width, original_height = img.size

                # Проверка, нужно ли изменять размер
                if original_width > max_width or original_height > max_height:
                    ratio = min(max_width / original_width, max_height / original_height)
                    new_size = (int(original_width * ratio), int(original_height * ratio))

                    # Универсально определяем метод ресемплинга
                    try:
                        resample = Image.Resampling.LANCZOS  # Новые версии Pillow
                    except AttributeError:
                        resample = Image.LANCZOS  # Старые версии Pillow

                    img = img.resize(new_size, resample)

                # Сохраняем в буфер
                buffer = io.BytesIO()
                ext = file_path.lower().split('.')[-1]
                format = "JPEG" if ext in ("jpg", "jpeg") else "PNG"
                img.save(buffer, format=format)
                buffer.seek(0)

                filename = file_path.split("/")[-1]
                mime_type = "image/jpeg" if format == "JPEG" else "image/png"
                files = {"file": (filename, buffer, mime_type)}
                headers = {"Authorization": f"Bearer {self.token}"}

                response = requests.post("http://127.0.0.1:8000/auth/avatar", files=files, headers=headers)

                if response.status_code == 200:
                    self.load_avatar()
                    QMessageBox.information(self, "Успех", "Аватар обновлён")
                else:
                    QMessageBox.warning(self, "Ошибка", "Не удалось загрузить аватар")

            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки:\n{e}")

    def set_profile_info(self):
        profile = self.profile_data
        text = f"""
<b>ФИО:</b> {profile.get('full_name', '—')}<br>
<b>Должность:</b> {profile.get('position', '—')}<br><br>
<b>Опыт:</b><br>{profile.get('experience', '—')}<br><br>
<b>Компетенции:</b><br>{profile.get('skills', '—')}<br><br>
<b>Достижения:</b><br>{profile.get('achievements', '—')}
"""
        self.profile_text.setHtml(text)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    viewer = FormViewer()
    viewer.show()
    sys.exit(app.exec_())
