import re
from docx import Document

doc = Document("шаблон.docx")

data = {
    "1": "Иван Иванов",
    "2": "25 лет"
}

pattern = re.compile(r"поле_(\d+)")

def replace_in_paragraph(paragraph):
    # Собираем весь текст абзаца
    full_text = ''.join(run.text for run in paragraph.runs)
    # Проверяем, есть ли что-то для замены
    matches = pattern.findall(full_text)
    if not matches:
        return
    for match in matches:
        if match in data:
            placeholder = f"Поле для заполнения_{match}"
            full_text = full_text.replace(placeholder, data[match])
    # Очищаем и перезаписываем абзац
    for run in paragraph.runs:
        run.text = ''
    paragraph.runs[0].text = full_text  # пишем всё в первый run

# Обычные абзацы
for para in doc.paragraphs:
    replace_in_paragraph(para)

# В таблицах
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_in_paragraph(para)

doc.save("заполненный_документ.docx")



